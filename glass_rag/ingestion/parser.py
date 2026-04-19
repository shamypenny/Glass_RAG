"""Document parsing and lightweight structured extraction."""

from __future__ import annotations

import base64
import json
import mimetypes
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field

from glass_rag.schemas import DocumentMetadata, ParsedDocument, StructuredSection


class ExtractionPayload(BaseModel):
    """LLM-extracted structured entities for domain-specific metadata enrichment."""

    material_names: list[str] = Field(default_factory=list)
    formulas: list[str] = Field(default_factory=list)
    parameters: list[str] = Field(default_factory=list)
    test_results: list[str] = Field(default_factory=list)
    process_steps: list[str] = Field(default_factory=list)
    keywords: list[str] = Field(default_factory=list)
    summary: str = ""


class DocumentParser:
    """Parses files from multiple formats into normalized sections."""

    def __init__(self, parser_llm, extractor_llm, title_max_length: int) -> None:
        self.parser_llm = parser_llm
        self.extractor_llm = extractor_llm
        self.title_max_length = title_max_length
        self.extract_prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You extract structured knowledge for a glass-material RAG system. "
                    "Identify material names, formulas, parameters, test results, process steps, and keywords.",
                ),
                ("human", "Title: {title}\n\nContent:\n{content}"),
            ]
        )

    def parse(
        self,
        file_path: str,
        metadata: DocumentMetadata,
        callbacks=None,
    ) -> ParsedDocument:
        suffix = Path(file_path).suffix.lower()
        if suffix in {".txt", ".md", ".csv", ".json"}:
            sections = self._parse_text_like(file_path)
        elif suffix == ".pdf":
            sections = self._parse_pdf(file_path)
        elif suffix == ".docx":
            sections = self._parse_docx(file_path)
        elif suffix in {".png", ".jpg", ".jpeg", ".bmp"}:
            sections = self._parse_image(file_path, callbacks=callbacks)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        extracted_entities = self._extract_entities(metadata.title, sections, callbacks=callbacks)
        return ParsedDocument(
            file_path=file_path,
            metadata=metadata,
            sections=sections,
            extracted_entities=extracted_entities,
        )

    def _parse_text_like(self, file_path: str) -> list[StructuredSection]:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
        return [StructuredSection(content=block) for block in blocks]

    def _parse_pdf(self, file_path: str) -> list[StructuredSection]:
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        sections: list[StructuredSection] = []
        for page in pages:
            text = page.page_content.strip()
            if not text:
                continue
            sections.append(
                StructuredSection(
                    content=text,
                    page_number=page.metadata.get("page", 0) + 1,
                )
            )
        return sections

    def _parse_docx(self, file_path: str) -> list[StructuredSection]:
        doc = DocxDocument(file_path)
        sections: list[StructuredSection] = []
        current_heading: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                sections.append(
                    StructuredSection(
                        heading=current_heading,
                        content="\n".join(buffer).strip(),
                    )
                )
                buffer.clear()

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and "Heading" in paragraph.style.name:
                flush()
                current_heading = text[: self.title_max_length]
            else:
                buffer.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows:
                flush()
                sections.append(
                    StructuredSection(
                        heading=current_heading,
                        content="\n".join(rows),
                        section_type="table",
                    )
                )

        flush()
        return sections

    def _parse_image(self, file_path: str, callbacks=None) -> list[StructuredSection]:
        if self.parser_llm is None:
            raise ValueError("Image OCR requires a configured multimodal parser model.")

        mime_type = mimetypes.guess_type(file_path)[0] or "image/jpeg"
        data = base64.b64encode(Path(file_path).read_bytes()).decode("utf-8")
        message = HumanMessage(
            content=[
                {
                    "type": "text",
                    "text": "Run OCR on this image and preserve titles, paragraphs, tables, and key numeric values.",
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{data}"},
                },
            ]
        )
        response = self.parser_llm.invoke([message], config={"callbacks": callbacks or []})
        content = response.content if isinstance(response.content, str) else json.dumps(response.content, ensure_ascii=False)
        return [StructuredSection(content=content)]

    def _extract_entities(self, title: str, sections: list[StructuredSection], callbacks=None) -> dict:
        if self.extractor_llm is None or not sections:
            return {}

        text = "\n\n".join(section.content for section in sections)[:6000]
        chain = self.extract_prompt | self.extractor_llm.with_structured_output(ExtractionPayload)
        try:
            result = chain.invoke(
                {"title": title, "content": text},
                config={"callbacks": callbacks or []},
            )
            return result.model_dump()
        except Exception:
            return {}

